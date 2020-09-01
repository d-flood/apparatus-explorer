import ctypes
import os
import platform
import re
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog as fd
from lxml import etree as ET


def disable_enable_buttons(current, app):
    # set verse nav buttons
    if current.getprevious().tag != "ab":
        prev_vrs_btn.config(state=DISABLED)
    else:
        prev_vrs_btn.config(state=NORMAL)
    if current.getnext() == None:
        next_vrs_btn.config(state=DISABLED)
    else:
        next_vrs_btn.config(state=NORMAL)
    # set index prev button
    app_prev = app.getprevious()
    if app_prev == None:
        prev_index_button.config(state=DISABLED)
    elif app_prev.tag == "seg":
        app_prev = app_prev.getprevious()
    if app_prev == None:
        prev_index_button.config(state=DISABLED)
    else:
        prev_index_button.config(state=NORMAL)
    # set index next button
    app_next = app.getnext()
    if app_next == None:
        next_index_button.config(state=DISABLED)
    elif app_next.tag == "seg":
        app_next = app_next.getnext()
    if app_next == None:
        next_index_button.config(state=DISABLED)
    else:
        next_index_button.config(state=NORMAL)

def update_index(current, app):
    verse = re.sub("-APP", "", current.get('verse'))
    apps = current.findall('app')
    global i_from, i_to
    if app == None:
        i_from = 0
        i_to = -1
    elif app.tag == "app":
        i_from = int(app.get('from'))
        i_to = int(app.get('to'))
    for widget in temp_frame.winfo_children():
        widget.destroy()
    temp_frame.pack(side=LEFT)
    # find and display all variation indexes
    if apps == []:
        app_label = Label(index_frame, text="No Variation Units",
                    font=("Times", "12"))
        app_label.pack(side=LEFT, padx=10)
    else:
        index = app.get('from')+" - "+app.get('to')
        for app_unit in apps:
            app_index = f"{app_unit.get('from')} - {app_unit.get('to')}"
            app_label = Label(temp_frame, text=app_index, font=("Times", "12"))
            app_label.pack(side=LEFT, padx=15)
            if app_index == index:
                app_label.config(bg="spring green")
    ref_entry.delete(0, END)
    ref_entry.insert(0, verse)
    get_basetext()
    get_arcs(app)

def refill(current):
    verse = re.sub("-APP", "", current.get('verse'))
    global app, i_from, i_to, tree
    app = current.find('app')
    apps = current.findall('app')
    for widget in temp_frame.winfo_children():
        widget.destroy()
    temp_frame.pack(side=LEFT)
    # find and display all variation indexes
    if apps == []:
        app_label = Label(
            index_frame, text="No Variation Units", font=("Times", "12"))
        app_label.pack(side=LEFT, padx=10)
    else:
        index = app.get('from')+" - "+app.get('to')
        for app_unit in apps:
            app_index = f"{app_unit.get('from')} - {app_unit.get('to')}"
            app_label = Label(temp_frame, text=app_index, font=("Times", "12"))
            app_label.pack(side=LEFT, padx=15)
            if app_index == index:
                app_label.config(bg="spring green")
        get_arcs(app)
    ref_entry.delete(0, END)
    ref_entry.insert(0, verse)
    get_basetext()

def updater(change, direction):
    global tree, app
    root = tree.getroot()
    ab = f"{ref_entry.get()}-APP"
    if change == "verse":
        current = root.find(f".//ab[@verse='{ab}']")
        if direction == "next":
            current = current.getnext()
            refill(current)
        elif direction == "prev":
            current = current.getprevious()
            refill(current)
        elif direction == "None":
            refill(current)
    elif change == "index":
        current = root.find(f".//ab[@verse='{ab}']")
        if direction == "next":
            app = app.getnext()
            if app.tag == "seg":
                app = app.getnext()
        if direction == "prev":
            app = app.getprevious()
            if app.tag == "seg":
                app = app.getprevious()
    elif change == "initial startup":
        if ref_entry.get() == "":
            current = root.find("ab")
            app = current.find('app')
        else:
            current = root.find(f".//ab[@verse='{ab}']")
            refill(current)
    update_index(current, app)
    disable_enable_buttons(current, app)

def browse():
    xml_dir = fd.askopenfilename(
        initialdir=f"{main_dir}/Collation/open-cbgm/Collations")
    xml_dir_entry.insert(0, xml_dir)

def get_arcs(app):
    global tree
    if app == None:
        pass
    else:
        rdgs = app.findall('rdg')
        for widget in rdgs_frame.winfo_children():
            widget.destroy()
        for rdg in rdgs:
            if rdg.text:
                greek_text = rdg.text
            else:
                greek_text = rdg.get('type')
            rdg_label = Label(
                rdgs_frame, 
                text=f"{rdg.get('n')}  {greek_text}\t{rdg.get('wit')}", 
                font=("Times", "12"), wraplength=1000)
            rdg_label.pack(side=TOP, anchor=W)
        arcs = app.findall('note/graph/arc')
        for widget in gen_frame.winfo_children():
            widget.destroy()
        if arcs == []:
            arc_label = Label(gen_frame, text="No Arcs Present",
                        font=("Times", "12"))
            arc_label.pack(side=TOP)
        else:
            for arc in arcs:
                arc_label = Label(
                    gen_frame, 
                    text=f"{arc.get('from')} --> {arc.get('to')}", 
                    font=("Times", "12"))
                arc_label.pack(side=TOP)

def load_xml():
    global tree
    cx_fname = xml_dir_entry.get()
    if cx_fname == "":
        messagebox.showinfo(title="Uh-oh", 
        message="File path field is blank. Type in the file path\
             (location) or click 'Browse'")
    else:
        with open(cx_fname, 'r', encoding='utf-8') as file:
            tree = file.read()
        tree = re.sub("xml:id", "verse", tree)
        tree = re.sub("<TEI xmlns=\"http://www.tei-c.org/ns/1.0\">", "<TEI>", tree)
        tree = re.sub("<?xml version='1.0' encoding='UTF-8'?>", "", tree)
        with open(cx_fname, 'w', encoding='utf-8') as file:
            file.write(tree)
        parser = ET.XMLParser(remove_blank_text=True)
        tree = ET.parse(cx_fname, parser)
        updater("initial startup", "None")

def save_exit():
    if tree == None:
        pass
    else:
        write_new_xml()
        print("XML File Saved")
    main.destroy()

def prev_btn_cmd():
    updater("verse", "prev")

def next_btn_cmd():
    updater("verse", "next")

def prev_index_cmd():
    updater("index", "prev")

def next_index_cmd():
    updater("index", "next")

def get_basetext():
    ref = str(ref_entry.get())
    ref = ref.split('.')
    chp = ref[0]
    vrs = ref[1]
    rp_fname = f"{main_dir}/mss/RP/Plain Text Transcriptions/RP_{chp}.txt"
    # Get verse from basetext
    with open(rp_fname, 'r', encoding='utf-8') as file:
        basetext = file.read()
        pass
    basetext = re.search(vrs + r"(.+)", basetext)
    basetext = re.sub(vrs, "", basetext.group(0))
    basetext = basetext.strip().split()
    index_list = []
    global i_from, i_to, index, tree
    for i in range(i_from, i_to+1):
        if i % 2 == 0:
            index_list.append(str(i))
    index = []
    count = 2
    for i in range(len(basetext)):
        count_str = str(count)
        index.append(count_str)
        count += 2
    for widget in basetext_frame.winfo_children():
        widget.destroy()
    for ind, word, in zip(index, basetext):
        index_label = Label(basetext_frame, text=ind+"\n"+word,
                        font=("Times", "12"))
        index_label.pack(side=LEFT, padx=5)
        if ind in index_list:
            index_label.config(bg="spring green")

def write_new_xml():
    global tree
    tree.write(xml_dir_entry.get(), encoding="utf-8", pretty_print=True)
    with open(xml_dir_entry.get(), 'r', encoding='utf-8') as file:
        new_file = file.read()
    new_file = re.sub("verse", "xml:id", new_file)
    new_file = re.sub(
        "<TEI>", 
        "<?xml version='1.0' encoding='UTF-8'?>\n<TEI xmlns=\"http://www.tei-c.org/ns/1.0\">", 
        new_file)
    with open(xml_dir_entry.get(), 'w', encoding='utf-8') as file:
        file.write(new_file)

def add_arc():
    if arc_entry_1.get() == "" or arc_entry_2.get() == "":
        messagebox.showinfo(title="Uh-oh", 
        message="Remember to fill both boxes in the 'Edit Relationships' section")
    else:
        global app, tree
        arc = ET.Element("arc")
        arc.set('from', arc_entry_1.get())
        arc.set('to', arc_entry_2.get())
        arc_parent = app.find("note/graph")
        arc_parent.append(arc)
        arcs = app.findall('note/graph/arc')
        for widget in gen_frame.winfo_children():
            widget.destroy()
        if arcs == []:
            arc_label = Label(gen_frame, text="No Arcs Present",
                        font=("Times", "12"))
            arc_label.pack(side=TOP)
        else:
            for item in arcs:
                arc_label = Label(gen_frame,
                            text=f"{item.get('from')} --> {item.get('to')}",
                            font=("Times", "12"))
                arc_label.pack(side=TOP)
        write_new_xml()
    
def delete_arc():
    if arc_entry_1.get() == "" or arc_entry_2.get() == "":
        messagebox.showinfo(title="Uh-oh", 
        message="Remember to fill both boxes in the 'Edit Relationships' section")
    else:
        global app
        arcs = app.findall('note/graph/arc')
        for arc in arcs:
            if arc.get('from') == arc_entry_1.get() and arc.get('to') == arc_entry_2.get():
                arc.getparent().remove(arc)
        for widget in gen_frame.winfo_children():
            widget.destroy()
        arcs = app.findall('note/graph/arc')
        if arcs == []:
            arc_label = Label(gen_frame, text="No Arcs Present",
                font=("Times", "12"))
            arc_label.pack(side=TOP)
        else:
            for arc in arcs:
                arc_label = Label(gen_frame,
                    text=f"{arc.get('from')} --> {arc.get('to')}",
                    font=("Times", "12"))
                arc_label.pack(side=TOP)
        write_new_xml()

def export_docx():
    from docx import Document
    global tree
    root = tree.getroot()
    document = Document(f"{main_dir}/template.docx")
    document.add_heading('Critical Apparatus\n', 0)

    ab_elements = root.findall("ab")
    for ab in ab_elements:
        apps = ab.findall('app')
        verse = ab.get("verse")
        verse = re.sub("-APP", "", verse)
        full_verse = re.sub("R", "\nRomans ", verse)
        full_verse = re.sub(r"\.", ":", full_verse)

        document.add_heading(full_verse, level=1)

        ref = re.sub(r"\.", " ", verse)
        ref = ref.split()
        chp = ref[0]
        vrs = ref[1]
        rp_fname = f"{main_dir}/mss/RP/Plain Text Transcriptions/RP_{chp}.txt"
        # Get RP verse for display and also the ECM style index nums
        with open(rp_fname, 'r', encoding='utf-8') as file:
            basetext = file.readlines()
            pass
        for line in basetext:
            if line.startswith(vrs):
                line = re.sub(vrs, "", line)
                basetext = line.strip().split()
        index = []
        count = 2

        for i in range(len(basetext)):
            count_str = str(count)
            index.append(count_str)
            count += 2

        verse_length = len(basetext)

        cell = 0
        if verse_length <= 15:
            table = document.add_table(rows=1, cols=verse_length)
            row_cells = table.add_row().cells
            for x, y in zip(index, basetext):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

        elif verse_length <= 30:
            # index_a = index[:15]
            # basetext_a = basetext[:15]

            table = document.add_table(rows=1, cols=15)
            row_cells = table.add_row().cells

            for x, y in zip(index[:15], basetext[:15]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1
            cell = 0
            row_cells = table.add_row().cells
            for x, y in zip(index[15:], basetext[15:]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

        else:
            table = document.add_table(rows=1, cols=15)
            row_cells = table.add_row().cells

            for x, y in zip(index[:15], basetext[:15]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

            cell = 0
            row_cells = table.add_row().cells
            for x, y in zip(index[15:30], basetext[15:30]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1
            
            cell = 0
            row_cells = table.add_row().cells
            for x, y in zip(index[30:], basetext[30:]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

        for app in apps:
            p = document.add_paragraph("\n"+app.get('from')+"–"+app.get('to')).underline = True
            rdgs = app.findall('rdg')
            for rdg in rdgs:
                if rdg.text:
                    greek_text = rdg.text
                    p = document.add_paragraph(f"Reading {rdg.get('n')}: ")
                    p.add_run(greek_text).bold = True
                    p.add_run(f"\t\t{rdg.get('wit')}")
                else:
                    greek_text = rdg.get('type')
                    p = document.add_paragraph(f"Reading {rdg.get('n')}: ")
                    p.add_run(greek_text).bold = True
                    p.add_run(f"\t\t{rdg.get('wit')}")

    document.save("apparatus.docx")
    print("Apparatus exported as docx")

"""GUI Startup"""
tree = None
main_dir = os.getcwd()
main_dir = re.sub(r"\\", "/", main_dir)

# tkinter seems to have a scaling issue on high-resolution screens on Windows.
# The following seems to solve the issue. If GUI is blurry, try changing
# the SetProcessDpiAwareness argument to a different number between 0 and 3.
my_os = platform.system()
if my_os == "Windows":
    ctypes.windll.shcore.SetProcessDpiAwareness(1)

main = Tk()
main.title("Apparatus Explorer v.03")
main.iconbitmap(main_dir+"/console.ico")

# FRAMES

xml_dir_frame  = LabelFrame(main, text="")
xml_dir_frame.grid(row=8, column=0, columnspan=3, pady=10, padx=10)

ref_frame = LabelFrame(main, text="")
ref_frame.grid(row=0, column=0, columnspan=3, pady=10, padx=10)

index_frame = LabelFrame(main, text="")
index_frame.grid(row=1, column=0, columnspan=3, pady=10, padx=10)

temp_frame = LabelFrame(index_frame) #packed after prev_index_btn

basetext_frame = LabelFrame(main, text="")
basetext_frame.grid(row=2, column=0, columnspan=3, pady=10, padx=10)

rdgs_frame = LabelFrame(main, text="Readings", font=("Times", "12"))
rdgs_frame.grid(row=4, column=0, columnspan=3, pady=10, padx=10)

gen_frame = LabelFrame(main, text="Genealogical Relationships", 
            font=("Times", "12"))
gen_frame.grid(row=5, column=0, pady=10, padx=10)

edit_arc_frame = LabelFrame(main, text="Edit Relationships",
                font=("Times", "12"))
edit_arc_frame.grid(row=5, column=1, padx=10, pady=10)

# LABELS
xml_dir_label = Label(xml_dir_frame, text="XML Collation File Path: ",
            font=("Times", "12"))
xml_dir_label.grid(row=0, column=0)

arc_label_1 = Label(edit_arc_frame, text="-->", font=("Times", "12"))
arc_label_1.grid(row=0, column=2, padx=10)

# BUTTONS

load_xml_button = Button(xml_dir_frame, 
    text="Load/Refresh XML File", 
    font=("Times", "12"), 
    command=load_xml)
load_xml_button.grid(row=0, column=3, padx=10)

browse_button = Button(xml_dir_frame, text="Browse", font=("Times", "12"),
                command=browse)
browse_button.grid(row=0, column=2, padx=10)

save_xml_btn = Button(xml_dir_frame, 
    text="Save and Exit", 
    font=("Times", "12"), 
    command=save_exit)
save_xml_btn.grid(row=0, column=4, padx=10)

export_button = Button(xml_dir_frame, text="Export docx",
        font=("Times", "12"), command=export_docx)
export_button.grid(row=0, column=5, padx=10)

prev_vrs_btn = Button(ref_frame, text="<Prev", font=("Times", "12"),
                command=prev_btn_cmd, state=DISABLED)
prev_vrs_btn.grid(row=0, column=0)

next_vrs_btn = Button(ref_frame, text="Next>", font=("Times", "12"),
                command=next_btn_cmd, state=DISABLED)
next_vrs_btn.grid(row=0, column=2)

prev_index_button = Button(index_frame, text="<Prev", font=("Times", "12"),
                    command=prev_index_cmd, state=DISABLED)
prev_index_button.pack(side=LEFT, padx=10)
temp_frame.pack(side=LEFT)

next_index_button = Button(index_frame, text="Next>", font=("Times", "12"),
                    command=next_index_cmd, state=DISABLED)
next_index_button.pack(side=RIGHT, padx=10)

update_arc_button = Button(edit_arc_frame, text="Add", font=("Times", "12"),
                    command=add_arc)
update_arc_button.grid(row=0, column=4, padx=10)

delete_arc_button = Button(edit_arc_frame, text="Delete", 
                    font=("Times", "12"), command=delete_arc)
delete_arc_button.grid(row=0, column=5, padx=10)

# ENTRY WIDGETS

arc_entry_1 = Entry(edit_arc_frame, width=3, font=("Times", "12"))
arc_entry_1.grid(row=0, column=1, padx=10)

arc_entry_2 = Entry(edit_arc_frame, width=3, font=("Times", "12"))
arc_entry_2.grid(row=0, column=3, padx=10)

xml_dir_entry = Entry(xml_dir_frame, width=50, font=("Times", "12"))
xml_dir_entry.grid(row=0, column=1)

ref_entry = Entry(ref_frame, width=30, font=("Times", "12"))
ref_entry.grid(row=0, column=1, padx=50)


main.mainloop()