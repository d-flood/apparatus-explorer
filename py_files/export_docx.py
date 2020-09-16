import datetime
import re
from lxml import etree as ET
from py_files.get_basetext import get_basetext

def export_docx(tree, main_dir):
    try:
        from docx import Document
    except:
        return 'could not import python-docx.\nMake sure it is installed: https://pypi.org/project/python-docx/'
    root = tree.getroot()
    document = Document(f"{main_dir}/files/template.docx")
    document.add_heading('Critical Apparatus\n', 0)

    ab_elements = root.findall("ab")
    for ab in ab_elements:
        apps = ab.findall('app')
        verse = ab.get("verse")
        verse = re.sub("-APP", "", verse)
        full_verse = re.sub("Rm", "\nRomans ", verse)
        full_verse = re.sub(r"\.", ":", full_verse)

        document.add_heading(full_verse, level=1)

        basetext = get_basetext(verse, main_dir)
        index = []
        count = 2

        for i in range(len(basetext)):
            count_str = str(count)
            index.append(count_str)
            count += 2

        verse_length = len(basetext)

        cell = 0
        if verse_length <= 15:
            table = document.add_table(rows=1, cols=verse_length)
            row_cells = table.add_row().cells
            for x, y in zip(index, basetext):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

        elif verse_length <= 30:

            table = document.add_table(rows=1, cols=15)
            row_cells = table.add_row().cells

            for x, y in zip(index[:15], basetext[:15]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1
            cell = 0
            row_cells = table.add_row().cells
            for x, y in zip(index[15:], basetext[15:]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

        else:
            table = document.add_table(rows=1, cols=15)
            row_cells = table.add_row().cells

            for x, y in zip(index[:15], basetext[:15]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

            cell = 0
            row_cells = table.add_row().cells
            for x, y in zip(index[15:30], basetext[15:30]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1
            
            cell = 0
            row_cells = table.add_row().cells
            for x, y in zip(index[30:], basetext[30:]):
                row_cells[cell].text = f"{x}\n{y}"
                cell += 1

        for app in apps:
            if app.get('from') == app.get('to'):
                p = document.add_paragraph("\n"+app.get('from')).underline = True
            else:
                p = document.add_paragraph("\n"+app.get('from')+"–"+app.get('to')).underline = True
            rdgs = app.findall('rdg')
            for rdg in rdgs:
                if rdg.text:
                    greek_text = rdg.text
                    p = document.add_paragraph(f"Reading {rdg.get('n')}: ")
                    p.add_run(greek_text).bold = True
                    p.add_run(f"\t\t{rdg.get('wit')}")
                else:
                    greek_text = rdg.get('type')
                    p = document.add_paragraph(f"Reading {rdg.get('n')}: ")
                    p.add_run(greek_text).bold = True
                    p.add_run(f"\t\t{rdg.get('wit')}")

    now = datetime.datetime.now()
    now_date = now.strftime('%x')
    now_date = now_date.replace('/', '-')
    now_time = now.strftime('%X')
    now_time = now_time.replace(':', '-')
    datetime_to_append = f"{now_date}_{now_time}"
    document.save(f"{main_dir}/exported/apparatus--{datetime_to_append}.docx")
    return 'Apparatus converted to docx and saved to /exported'